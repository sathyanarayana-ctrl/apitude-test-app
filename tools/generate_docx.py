"""Generate Aptitude Test App documentation as Word (.docx)."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "Aptitude_Test_App_Guide.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(f"{bold_prefix}: ")
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def build_document():
    doc = Document()

    title = doc.add_heading("Aptitude Test App - User Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Cross-platform Flutter application for practicing aptitude questions "
        "across all major competitive exam topics. Built for Android, iOS, and Web."
    )
    doc.add_paragraph("GitHub: https://github.com/sathyanarayana-ctrl/apitude-test-app")
    doc.add_paragraph("Author: Satyanarayana")

    add_heading(doc, "1. App Overview", 1)
    doc.add_paragraph(
        "The Aptitude Test App helps users prepare for competitive exams with "
        "multiple-choice questions (MCQs), timed mock tests, category-wise practice, "
        "and full test papers. After each test, users receive a score, grade, and "
        "detailed answer review with explanations."
    )

    add_heading(doc, "2. Main Features", 1)
    features = [
        ("13 aptitude categories", "Quantitative, Logical, Verbal, Data Interpretation, and more"),
        ("Practice by category", "Focus on one topic at a time"),
        ("Test Papers", "13 built-in test papers plus Google Sheets import"),
        ("Full Mock Test", "20 mixed questions with a 20-minute timer"),
        ("Quick Practice", "10 random questions with no timer"),
        ("MCQ format", "A / B / C / D answer options"),
        ("Question navigator", "Jump to any question during a test"),
        ("Results and review", "Score, grade, and explanations for every answer"),
    ]
    for name, desc in features:
        add_bullet(doc, desc, bold_prefix=name)

    add_heading(doc, "3. How the App Works", 1)

    add_heading(doc, "3.1 Home Screen", 2)
    doc.add_paragraph("When the app opens, the user sees four main options:")
    home_options = [
        "Practice by Category - Pick one aptitude topic",
        "Test Papers - Open full test papers (13 built-in + Google Sheet)",
        "Full Mock Test - 20 random questions with 20-minute timer",
        "Quick Practice - 10 random questions with no timer",
    ]
    for item in home_options:
        add_bullet(doc, item)

    add_heading(doc, "3.2 Category and Test Papers Screens", 2)
    doc.add_paragraph(
        "Practice by Category shows all 13 aptitude types. Tapping a category "
        "starts a quiz with questions for that topic only."
    )
    doc.add_paragraph(
        "Test Papers shows 13 full test papers (5 questions each), one per category. "
        "It also loads questions from the connected Google Sheet when shared publicly. "
        "Use the refresh button to reload from Google Sheets."
    )

    add_heading(doc, "3.3 Quiz Screen", 2)
    doc.add_paragraph("During a test, the user can:")
    quiz_actions = [
        "Read the question at the top",
        "Tap an option (A, B, C, or D) to select an answer",
        "Use Previous and Next buttons to move between questions",
        "Jump to any question using numbered boxes",
        "See the countdown timer (for timed tests)",
        "Submit on the last question or when time runs out",
    ]
    for item in quiz_actions:
        add_bullet(doc, item)

    add_heading(doc, "3.4 Results Screen", 2)
    doc.add_paragraph("After submission, the app shows:")
    result_items = [
        "Score percentage and grade (Excellent, Good, Needs Practice, etc.)",
        "Correct, wrong, and skipped question counts",
        "Answer review for every question with explanations",
        "Back to Home button to start another test",
    ]
    for item in result_items:
        add_bullet(doc, item)

    add_heading(doc, "4. Question Categories", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Topics Covered"
    categories = [
        ("Quantitative Aptitude", "Percentages, ratios, speed, algebra"),
        ("Logical Reasoning", "Statements, assumptions, conclusions"),
        ("Verbal Ability", "Synonyms, antonyms, grammar"),
        ("Data Interpretation", "Charts, tables, averages"),
        ("Analytical Reasoning", "Seating, ranking, direction"),
        ("Non-Verbal Reasoning", "Patterns, mirror images, figures"),
        ("Syllogism", "Logical conclusions from statements"),
        ("Blood Relations", "Family tree puzzles"),
        ("Coding & Decoding", "Letter and number codes"),
        ("Clock & Calendar", "Angles, days, dates"),
        ("Series", "Number and letter series"),
        ("Puzzles", "Brain teasers and logic puzzles"),
        ("General Aptitude", "SI, HCF, work & time, mixed"),
    ]
    for cat, topics in categories:
        row = table.add_row().cells
        row[0].text = cat
        row[1].text = topics

    add_heading(doc, "5. Question Data Sources", 1)
    sources = [
        "questions_repository.dart - 35 built-in practice questions",
        "assets/test_papers/all_papers.json - 13 test papers with 65 questions",
        "Google Sheets - Live import from shared spreadsheet",
    ]
    for item in sources:
        add_bullet(doc, item)

    doc.add_paragraph(
        "Google Sheet URL: "
        "https://docs.google.com/spreadsheets/d/1wqayZDw8mLilWHrrSHVtOO20gr1PyfTCITlAMBd5o-o/edit"
    )

    add_heading(doc, "6. Google Sheet Format", 1)
    doc.add_paragraph("Use these columns in row 1 (header row):")
    doc.add_paragraph(
        "id | type | question | option_a | option_b | option_c | option_d | "
        "correct_index | explanation | difficulty"
    )
    doc.add_paragraph("Share settings: Anyone with the link -> Viewer")

    add_heading(doc, "7. Project Structure", 1)
    structure = [
        "lib/main.dart - App entry point",
        "lib/screens/home_screen.dart - Main menu",
        "lib/screens/category_screen.dart - Category list",
        "lib/screens/test_papers_screen.dart - Test papers list",
        "lib/screens/quiz_screen.dart - Active quiz UI",
        "lib/screens/result_screen.dart - Score and review",
        "lib/providers/quiz_session.dart - Quiz logic, timer, scoring",
        "lib/data/questions_repository.dart - Practice questions",
        "lib/data/test_papers_repository.dart - Test paper loader",
        "lib/data/test_papers_config.dart - Google Sheet configuration",
        "lib/services/google_sheets_loader.dart - Google Sheets fetcher",
        "assets/test_papers/all_papers.json - Built-in test paper data",
    ]
    for item in structure:
        add_bullet(doc, item)

    add_heading(doc, "8. Installation and Run Instructions", 1)
    steps = [
        "Install Flutter SDK (3.2 or later)",
        "Clone: git clone https://github.com/sathyanarayana-ctrl/apitude-test-app.git",
        "cd apitude-test-app",
        "flutter create .",
        "flutter pub get",
        "flutter run",
        "For web: flutter run -d chrome",
        "For Android: connect device and run flutter run",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    add_heading(doc, "9. Example User Walkthrough", 1)
    walkthrough = [
        "Open the app and land on Home screen",
        "Tap Test Papers",
        "Select Quantitative Aptitude Test Paper 1",
        "Answer 5 questions using Next and Submit",
        "View score and read explanations on Results screen",
        "Tap Back to Home to start another test",
    ]
    for i, step in enumerate(walkthrough, 1):
        doc.add_paragraph(f"Step {i}: {step}")

    add_heading(doc, "10. Roadmap", 1)
    roadmap = [
        "Add more questions (100+ per category)",
        "User login and score history",
        "Hindi and regional language support",
        "Offline mode with local database",
        "Play Store and App Store release",
    ]
    for item in roadmap:
        add_bullet(doc, item)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build_document()
